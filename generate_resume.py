import os
import sys
import subprocess

# 1. Programmatically install python-docx if not present
try:
    import docx
except ImportError:
    print("Installing python-docx library...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_resume():
    doc = Document()
    
    # Page setup - 0.75 in margins for cleaner resume spacing
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Styles & Colors (matching portfolio's violet/indigo aesthetic)
    COLOR_PRIMARY = RGBColor(99, 102, 241)     # Indigo/Violet #6366f1
    COLOR_TEXT = RGBColor(40, 40, 40)          # Dark grey for body text
    COLOR_MUTED = RGBColor(110, 110, 110)      # Muted grey for dates

    # Helper function to style runs
    def style_run(run, font_name="Calibri", font_size=11, bold=False, italic=False, color=None):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    # Helper to add section headers with an elegant bottom border
    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text.upper())
        style_run(run, font_name="Calibri", font_size=12, bold=True, color=COLOR_PRIMARY)
        
        # XML to add bottom border/rule (horizontal line under header)
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                         r'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="A5B4FC"/>'
                         r'</w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)

    # Helper to add standard bullet list item with tight spacing
    def add_bullet(text_bold_prefix, text_body):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        
        if text_bold_prefix:
            run_bold = p.add_run(text_bold_prefix + ": ")
            style_run(run_bold, font_name="Calibri", font_size=10.5, bold=True, color=COLOR_TEXT)
            
        run_body = p.add_run(text_body)
        style_run(run_body, font_name="Calibri", font_size=10.5, color=COLOR_TEXT)

    # ─── HEADER SECTION ───────────────────────────────────────────
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.space_after = Pt(2)
    
    run_name = p_header.add_run("TISHA CHAWLANI\n")
    style_run(run_name, font_name="Calibri", font_size=24, bold=True, color=COLOR_PRIMARY)
    
    run_title = p_header.add_run("Software Engineer & AI Specialist\n")
    style_run(run_title, font_name="Calibri", font_size=13, bold=True, color=COLOR_TEXT)
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(12)
    
    contact_info = (
        "Email: tishachawlani95@gmail.com   |   Phone: +91-6261371431   |   "
        "LinkedIn: linkedin.com/in/tishachawlani   |   GitHub: github.com/Tisha717"
    )
    run_contact = p_contact.add_run(contact_info)
    style_run(run_contact, font_name="Calibri", font_size=10, color=COLOR_MUTED)

    # ─── SUMMARY ──────────────────────────────────────────────────
    add_section_header("Professional Summary")
    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_after = Pt(8)
    p_summary.paragraph_format.line_spacing = 1.15
    run_sum = p_summary.add_run(
        "Highly motivated and award-winning Software Engineer specializing in AI development, Multi-Agent systems, "
        "and DevOps. Proven track record of building intelligent, automated solutions (ranging from Agentic ITSM workflows "
        "to HR goal generation tools) and managing CI/CD platform operations for Fortune 500 clients. Recognized as a "
        "Google WE Scholar (Top 1%) and an IET India Scholarship National Finalist."
    )
    style_run(run_sum, font_name="Calibri", font_size=10.5, color=COLOR_TEXT)

    # ─── EDUCATION ────────────────────────────────────────────────
    add_section_header("Education")
    p_edu = doc.add_paragraph()
    p_edu.paragraph_format.space_before = Pt(4)
    p_edu.paragraph_format.space_after = Pt(2)
    
    run_degree = p_edu.add_run("B.Tech in Computer Science & Engineering (Specialization in Data Science)\n")
    style_run(run_degree, font_name="Calibri", font_size=11, bold=True, color=COLOR_TEXT)
    
    run_univ = p_edu.add_run("Amity University Chhattisgarh  |  Graduation: 2025\n")
    style_run(run_univ, font_name="Calibri", font_size=10, bold=False, color=COLOR_TEXT)
    
    run_grades = p_edu.add_run("CGPA: 9.1/10 (Highest Graduate CGPA, Academic Silver Medal recipient)")
    style_run(run_grades, font_name="Calibri", font_size=10, italic=True, color=COLOR_MUTED)

    # ─── EXPERIENCE ───────────────────────────────────────────────
    add_section_header("Professional Experience")
    
    # Job 1: Capgemini
    p_job1 = doc.add_paragraph()
    p_job1.paragraph_format.space_before = Pt(6)
    p_job1.paragraph_format.space_after = Pt(2)
    p_job1.paragraph_format.keep_with_next = True
    
    r1 = p_job1.add_run("Capgemini ")
    style_run(r1, font_name="Calibri", font_size=11, bold=True, color=COLOR_PRIMARY)
    r2 = p_job1.add_run("|  Software Engineer  ")
    style_run(r2, font_name="Calibri", font_size=11, bold=True, color=COLOR_TEXT)
    
    # Right-aligned date placeholder
    r3 = p_job1.add_run("\t\t\t\t\t\t\t\t\t\tOct 2025 - Present")
    style_run(r3, font_name="Calibri", font_size=10, bold=True, color=COLOR_MUTED)
    
    add_bullet("Starbucks Platform DevOps Team", 
               "Currently allocated to the Starbucks platform DevOps team. Responsibilities include managing enterprise-grade infrastructure, CI/CD pipelines, SPN rotations, Apigee API Gateways, Front Door, APIM configurations, and data loading/offboarding operations for high-priority platform services (DeepBrew, EDAP).")
    add_bullet("Banner Health ITSM Workflow POC", 
               "Developed an end-to-end agentic AI workflow on the Amplifier platform (ServiceNow AI Studio) for Tomcat service tickets. Used Python and scikit-learn for Pearson correlation and NLP frequency analysis (showing -0.81 correlation between assignment group and resolution time), guiding the design of a optimized, semi-autonomous 5-agent pipeline that automates ticket routing, resolution drafting, and closure.")
    add_bullet("GetSUCCESS Priority Autopilot", 
               "Engineered an internal multi-agent AI assistant using Streamlit, Microsoft AutoGen (3 specialized agents), and Llama 3.3 via Groq to auto-generate personalized, accurately weighted annual goals for HR systems, with session data persisted in a local SQLite database.")
    add_bullet("Smart Resume screening Platform", 
               "Built an end-to-end recruiter search and screening engine using FastAPI, Azure AI Search semantic ranking, spaCy skill extraction, and DeepSeek LLM for interview preparation workflows during the 3-month Capgemini training period.")
    add_bullet("L&D Nexus", 
               "Developed a full-stack AI-powered learning and development application utilizing FastAPI and React (Vite) with Recharts batch analytics and a LangChain, HuggingFace, and FAISS-based RAG pipeline to query organizational knowledge bases.")

    # Job 2: Globally Recruit
    p_job2 = doc.add_paragraph()
    p_job2.paragraph_format.space_before = Pt(6)
    p_job2.paragraph_format.space_after = Pt(2)
    p_job2.paragraph_format.keep_with_next = True
    
    r1 = p_job2.add_run("Globally Recruit ")
    style_run(r1, font_name="Calibri", font_size=11, bold=True, color=COLOR_PRIMARY)
    r2 = p_job2.add_run("|  University Relations Associate  ")
    style_run(r2, font_name="Calibri", font_size=11, bold=True, color=COLOR_TEXT)
    r3 = p_job2.add_run("\t\t\t\t\t\t\tApr 2025 - Sep 2025")
    style_run(r3, font_name="Calibri", font_size=10, bold=True, color=COLOR_MUTED)
    
    add_bullet(None, "Streamlined representative travel, operational engagement dashboards, and partnerships for Catholic University of America.")
    add_bullet(None, "Coordinated student acquisition and教育fairs engagements virtual sessions for Saint Louis University.")

    # Job 3: Excelerate
    p_job3 = doc.add_paragraph()
    p_job3.paragraph_format.space_before = Pt(6)
    p_job3.paragraph_format.space_after = Pt(2)
    p_job3.paragraph_format.keep_with_next = True
    
    r1 = p_job3.add_run("Excelerate ")
    style_run(r1, font_name="Calibri", font_size=11, bold=True, color=COLOR_PRIMARY)
    r2 = p_job3.add_run("|  AI Data Analyst & Prompt Researcher  ")
    style_run(r2, font_name="Calibri", font_size=11, bold=True, color=COLOR_TEXT)
    r3 = p_job3.add_run("\t\t\t\t\t\t\tFeb 2025 - Mar 2025")
    style_run(r3, font_name="Calibri", font_size=10, bold=True, color=COLOR_MUTED)
    
    add_bullet("Churn Analysis & Retention Modelling", 
               "Led a team of 12 to conduct EDA, hypothesis testing, and machine learning models (Random Forest, SVM, Decision Tree, Logistic Regression) on 8,500+ records, achieving a 99.7% model accuracy.")
    add_bullet("Prompt Engineering Roadmaps", 
               "Led a team of 6 researching LLM prompting strategies across Google Gemini, Hugging Face, and OpenAI, developing structured comparisons and integration blueprints.")

    # Other Job Mappings (Brief)
    p_other_jobs = doc.add_paragraph()
    p_other_jobs.paragraph_format.space_before = Pt(6)
    p_other_jobs.paragraph_format.space_after = Pt(2)
    
    r_oj = p_other_jobs.add_run("Additional Professional Experience:\n")
    style_run(r_oj, font_name="Calibri", font_size=10.5, bold=True, color=COLOR_TEXT)
    
    run_oj_details = (
        "• Math AI Trainer (Outlier, Sep 2024 - Nov 2024): Evaluated LLM mathematical reasoning, prompt optimization, and error analysis.\n"
        "• Analyst (AmPm Overseas, Jun 2024 - Nov 2024): Supported UK market expansion using CRM data segmentation.\n"
        "• Data Analyst Intern (NoQs Digital, May 2023 - Aug 2023): Automated workflows using Google Sheets API, reducing manual effort by 60%."
    )
    r_oj_det = p_other_jobs.add_run(run_oj_details)
    style_run(r_oj_det, font_name="Calibri", font_size=10, color=COLOR_TEXT)

    # ─── SKILLS ───────────────────────────────────────────────────
    add_section_header("Skills & Technologies")
    p_skills = doc.add_paragraph()
    p_skills.paragraph_format.space_after = Pt(6)
    p_skills.paragraph_format.line_spacing = 1.15
    
    s_pro = p_skills.add_run("Programming Languages: ")
    style_run(s_pro, font_name="Calibri", font_size=10.5, bold=True, color=COLOR_TEXT)
    s_pro_val = p_skills.add_run("Python, Java, SQL, MySQL\n")
    style_run(s_pro_val, font_name="Calibri", font_size=10.5, color=COLOR_TEXT)
    
    s_lib = p_skills.add_run("AI / Data Science: ")
    style_run(s_lib, font_name="Calibri", font_size=10.5, bold=True, color=COLOR_TEXT)
    s_lib_val = p_skills.add_run("Pandas, NumPy, scikit-learn, TensorFlow, OpenCV, Plotly, NLP, Prompt Engineering\n")
    style_run(s_lib_val, font_name="Calibri", font_size=10.5, color=COLOR_TEXT)
    
    s_frameworks = p_skills.add_run("Frameworks & Tools: ")
    style_run(s_frameworks, font_name="Calibri", font_size=10.5, bold=True, color=COLOR_TEXT)
    s_frameworks_val = p_skills.add_run("FastAPI, React, Microsoft AutoGen, Power BI, Tableau, Git\n")
    style_run(s_frameworks_val, font_name="Calibri", font_size=10.5, color=COLOR_TEXT)
    
    s_cloud = p_skills.add_run("Cloud & DevOps: ")
    style_run(s_cloud, font_name="Calibri", font_size=10.5, bold=True, color=COLOR_TEXT)
    s_cloud_val = p_skills.add_run("Azure DevOps, Azure AI Search, Azure Blob Storage, Apigee API Gateway, APIM, Front Door\n")
    style_run(s_cloud_val, font_name="Calibri", font_size=10.5, color=COLOR_TEXT)

    # ─── AWARDS ───────────────────────────────────────────────────
    add_section_header("Selected Honors & Awards")
    
    add_bullet("Google WE Scholar", 
               "Selected in the top 1% women in the nation for the TalentSprint & Google Women Engineers Program. Completed a rigorous 2-year training and mentorship program, participated in a 15-day workshop at IIIT Hyderabad, and visited the Google office.")
    add_bullet("IET India Scholarship Runner-Up", 
               "Top 10 National Finalist and Regional Runner-Up (2nd Prize, East & North East Zone) out of 65,000+ applicants. Presented the GreenGrow Hub sustainability solution at Calcutta University.")
    add_bullet("Harvard HPAIR Delegate", 
               "Selected as a delegate for the Harvard Project for Asian and International Relations (HPAIR) Conference 2023, focused on global leadership.")
    add_bullet("Shri Baljit Shastri Award & Silver Medal", 
               "Received the highest student honor of Amity University and the Convocation Silver Medal presented by the Hon'ble Governor of Chhattisgarh.")
    add_bullet("Smart India Hackathon", 
               "Winner of the university-level internal round (2023), representing Amity at the national SIH level with an AI chatbot for the Namami Gange Mission. Also represented the university as a participant at the national level in 2022.")
    add_bullet("Best & Star Performer", 
               "Awarded Best Performer (Prompt Engineering) and Star Performer across Excelerate internships in collaboration with Rochester Institute of Technology (RIT).")
    
    p_other_awards = doc.add_paragraph()
    p_other_awards.paragraph_format.space_before = Pt(4)
    p_other_awards.paragraph_format.space_after = Pt(2)
    r_oa = p_other_awards.add_run("Other Recognitions: ")
    style_run(r_oa, font_name="Calibri", font_size=10, bold=True, color=COLOR_TEXT)
    r_oa_val = p_other_awards.add_run(
        "AnitaB.org Advancing Inclusion Scholar (Grace Hopper Celebration India 2024)  |  "
        "Infosys Pragati (infosys Springboard Cohort 3 leadership invitee)  |  "
        "SheFi Scholar Season 12 (Web3 & DeFi)  |  "
        "Perspektywy Women in Tech Summit Attendee (2023)  |  "
        "G100 WEF & WICCI Invitee (2022)"
    )
    style_run(r_oa_val, font_name="Calibri", font_size=9.5, color=COLOR_MUTED)

    # ─── VOLUNTEERING ─────────────────────────────────────────────
    add_section_header("Community & Leadership")
    p_vol = doc.add_paragraph()
    p_vol.paragraph_format.space_before = Pt(4)
    p_vol.paragraph_format.space_after = Pt(2)
    
    r_vol = p_vol.add_run(
        "• Google Developer Student Clubs (GDSC) Lead (Jul 2022 - Aug 2023): Led core team of 10+, hosted 12+ tech workshops.\n"
        "• Newton School Coding Club (NSCC) President (Jul 2022 - Jul 2023): Coordinated programming contests for 100+ participants.\n"
        "• TFUG/ML Durg Core Member & Speaker: Delivered ML and data analysis sessions at ML Study Jams.\n"
        "• Women Techmakers Member & DevFest Raipur 2022 / Durg FOSS 2023 Organizing Member."
    )
    style_run(r_vol, font_name="Calibri", font_size=10, color=COLOR_TEXT)

    # Save Document
    filename = "Tisha_Chawlani_CV_updated.docx"
    doc.save(filename)
    print(f"Resume generated successfully: {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_resume()
